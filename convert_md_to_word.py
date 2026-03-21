#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Convert Markdown test case files to Word documents with table borders
"""
import os
import re
import markdown
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """Set cell borders"""
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), 'single')
            edge_el.set(qn('w:sz'), '4')
            edge_el.set(qn('w:space'), '0')
            edge_el.set(qn('w:color'), '000000')
            tcBorders.append(edge_el)
    
    tcPr.append(tcBorders)

def parse_markdown_file(md_path):
    """Parse markdown file and extract content"""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    return content

def create_word_doc_from_md(md_path, output_path):
    """Convert markdown test case file to Word document"""
    doc = Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # 标题处理
        if line.startswith('# '):
            heading_text = line[2:].strip()
            heading = doc.add_heading(heading_text, level=1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            i += 1
            continue
        
        if line.startswith('## '):
            heading_text = line[3:].strip()
            heading = doc.add_heading(heading_text, level=2)
            i += 1
            continue
        
        if line.startswith('### '):
            heading_text = line[4:].strip()
            heading = doc.add_heading(heading_text, level=3)
            i += 1
            continue
        
        # 处理表格
        if line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].rstrip())
                i += 1
            
            if table_lines:
                # 解析表格
                rows = []
                for table_line in table_lines:
                    cells = [cell.strip() for cell in table_line.split('|')[1:-1]]
                    rows.append(cells)
                
                if len(rows) >= 2:
                    # 跳过分隔符行
                    header = rows[0]
                    data_rows = rows[2:] if len(rows) > 2 else []
                    
                    if data_rows or (len(rows) > 1):
                        # 创建表格
                        table = doc.add_table(rows=1, cols=len(header))
                        table.style = 'Table Grid'
                        
                        # 设置表头
                        header_cells = table.rows[0].cells
                        for idx, header_text in enumerate(header):
                            header_cells[idx].text = header_text
                            # 设置表头样式
                            for paragraph in header_cells[idx].paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True
                                    run.font.size = Pt(11)
                        
                        # 添加数据行
                        for row_data in data_rows:
                            row_cells = table.add_row().cells
                            for idx, cell_text in enumerate(row_data):
                                if idx < len(row_cells):
                                    row_cells[idx].text = cell_text
                                    # 设置单元格文字大小
                                    for paragraph in row_cells[idx].paragraphs:
                                        for run in paragraph.runs:
                                            run.font.size = Pt(10)
                        
                        # 设置所有单元格边框
                        for row in table.rows:
                            for cell in row.cells:
                                set_cell_border(cell, top={}, left={}, bottom={}, right={}, insideH={}, insideV={})
            continue
        
        # 处理常规段落
        if line and not line.startswith('-'):
            if line.startswith('---'):
                doc.add_paragraph()
                i += 1
                continue
            
            # 添加段落
            if line:
                p = doc.add_paragraph(line)
                for run in p.runs:
                    run.font.size = Pt(10)
            i += 1
            continue
        
        i += 1
    
    doc.save(output_path)
    return output_path

def convert_single_file(index, filename):
    """Convert a single file"""
    input_dir = r'e:\IdeaProjects\flower-market\docs\test'
    output_dir = r'e:\IdeaProjects\flower-market\docs\test\Word'
    
    # 创建输出目录
    os.makedirs(output_dir, exist_ok=True)
    
    md_path = os.path.join(input_dir, filename)
    output_filename = filename.replace('.md', '.docx')
    output_path = os.path.join(output_dir, output_filename)
    
    if not os.path.exists(md_path):
        return None, f"文件不存在: {md_path}"
    
    try:
        result_path = create_word_doc_from_md(md_path, output_path)
        return result_path, "成功"
    except Exception as e:
        return None, str(e)

# 文件列表
FILES = [
    '功能覆盖表.md',
    '用户认证模块.md',
    '商品管理模块.md',
    '商品分类模块.md',
    '购物车模块.md',
    '订单管理模块.md',
    '优惠券模块.md',
    '商品评价模块.md',
    '商品收藏模块.md',
    '用户资料与余额模块.md',
    '收货地址模块.md',
    '签到奖励模块.md',
    '花卉知识模块.md',
    '商家管理模块.md',
    '管理员管理模块.md',
    '系统配置模块.md',
    '供应商管理模块.md',
    '商品溯源模块.md',
]

if __name__ == '__main__':
    print("程序已准备好，调用 convert_single_file(index, filename) 来处理单个文件")
    print("例如: convert_single_file(1, '用户认证模块.md')")
