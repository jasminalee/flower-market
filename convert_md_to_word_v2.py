#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Convert Markdown test case files to Word documents with table borders
"""
import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, color="000000"):
    """Set cell borders with proper formatting"""
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    for edge in ('top', 'left', 'bottom', 'right'):
        border_el = OxmlElement(f'w:{edge}')
        border_el.set(qn('w:val'), 'single')
        border_el.set(qn('w:sz'), '12')  # 边框大小
        border_el.set(qn('w:space'), '0')
        border_el.set(qn('w:color'), color)
        tcBorders.append(border_el)
    
    tcPr.append(tcBorders)

def parse_table_from_md(lines, start_idx):
    """Parse a markdown table starting from start_idx"""
    if not lines[start_idx].strip().startswith('|'):
        return None, start_idx
    
    table_lines = []
    i = start_idx
    
    while i < len(lines) and lines[i].strip().startswith('|'):
        line = lines[i].strip()
        if line:
            table_lines.append(line)
        i += 1
    
    if len(table_lines) < 2:
        return None, i
    
    # 解析表格数据
    rows = []
    for line in table_lines:
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        rows.append(cells)
    
    # 跳过分隔符行（第二行通常是 --- 分隔符）
    if len(rows) > 1 and all('-' in cell or cell == '' for cell in rows[1]):
        header = rows[0]
        data_rows = rows[2:]
    else:
        header = rows[0]
        data_rows = rows[1:]
    
    return {'header': header, 'data': data_rows}, i

def create_word_doc_from_md(md_path, output_path):
    """Convert markdown test case file to Word document"""
    doc = Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        
        # 处理 H1 标题
        if stripped.startswith('# '):
            text = stripped[2:].strip()
            heading = doc.add_heading(text, level=1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            i += 1
            continue
        
        # 处理 H2 标题
        if stripped.startswith('## '):
            text = stripped[3:].strip()
            heading = doc.add_heading(text, level=2)
            i += 1
            continue
        
        # 处理 H3 标题
        if stripped.startswith('### '):
            text = stripped[4:].strip()
            heading = doc.add_heading(text, level=3)
            i += 1
            continue
        
        # 处理表格
        if stripped.startswith('|'):
            table_data, next_idx = parse_table_from_md(lines, i)
            
            if table_data:
                header = table_data['header']
                data_rows = table_data['data']
                
                # 创建表格
                num_cols = len(header)
                num_rows = len(data_rows) + 1  # +1 for header
                
                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.style = 'Table Grid'
                
                # 设置表头
                header_cells = table.rows[0].cells
                for idx, header_text in enumerate(header):
                    cell = header_cells[idx]
                    cell.text = header_text
                    
                    # 格式化表头
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(10)
                    
                    set_cell_border(cell)
                
                # 添加数据行
                for row_idx, row_data in enumerate(data_rows):
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < len(row_cells):
                            cell = row_cells[col_idx]
                            cell.text = cell_text
                            
                            # 格式化单元格
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(9)
                                    # 左对齐
                                    if col_idx == 0:
                                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            
                            set_cell_border(cell)
            
            i = next_idx
            continue
        
        # 处理水平分隔线
        if stripped.startswith('---'):
            doc.add_paragraph()
            i += 1
            continue
        
        # 处理普通段落
        if stripped and not stripped.startswith('|'):
            p = doc.add_paragraph(stripped)
            for run in p.runs:
                run.font.size = Pt(10)
        
        i += 1
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
    
    doc.save(output_path)
    return output_path

def convert_single_file(filename):
    """Convert a single markdown file to Word"""
    input_dir = r'e:\IdeaProjects\flower-market\docs\test'
    output_dir = r'e:\IdeaProjects\flower-market\docs\test\Word'
    
    # 创建输出目录
    os.makedirs(output_dir, exist_ok=True)
    
    md_path = os.path.join(input_dir, filename)
    output_filename = filename.replace('.md', '.docx')
    output_path = os.path.join(output_dir, output_filename)
    
    if not os.path.exists(md_path):
        return False, f"错误: 文件不存在 - {md_path}"
    
    try:
        result_path = create_word_doc_from_md(md_path, output_path)
        return True, f"成功! 文件已保存到: {output_path}"
    except Exception as e:
        import traceback
        return False, f"错误: {str(e)}\n{traceback.format_exc()}"

# 文件列表（按顺序）
FILES = [
    ('1', '功能覆盖表.md'),
    ('2', '用户认证模块.md'),
    ('3', '商品管理模块.md'),
    ('4', '商品分类模块.md'),
    ('5', '购物车模块.md'),
    ('6', '订单管理模块.md'),
    ('7', '优惠券模块.md'),
    ('8', '商品评价模块.md'),
    ('9', '商品收藏模块.md'),
    ('10', '用户资料与余额模块.md'),
    ('11', '收货地址模块.md'),
    ('12', '签到奖励模块.md'),
    ('13', '花卉知识模块.md'),
    ('14', '商家管理模块.md'),
    ('15', '管理员管理模块.md'),
    ('16', '系统配置模块.md'),
    ('17', '供应商管理模块.md'),
    ('18', '商品溯源模块.md'),
]

if __name__ == '__main__':
    # 处理第一个文件
    success, message = convert_single_file('功能覆盖表.md')
    print(f"处理第1个文件: 功能覆盖表.md")
    print(f"结果: {message}")
    print(f"状态: {'成功' if success else '失败'}")
