from docx import Document

doc = Document(r'e:\IdeaProjects\flower-market\docs\test\Word\测试用例汇总.docx')
print("=== PARAGRAPHS ===")
for i, p in enumerate(doc.paragraphs[:40]):
    if p.text.strip():
        print(f"[P{i:02d}] [{p.style.name}] {p.text[:100]!r}")

print("\n=== TABLES ===")
for ti, tbl in enumerate(doc.tables):
    print(f"\n[Table {ti}] {len(tbl.rows)} rows x {len(tbl.columns)} cols")
    for ri, row in enumerate(tbl.rows[:4]):
        cells = [c.text[:40] for c in row.cells]
        print(f"  row{ri}: {cells}")
