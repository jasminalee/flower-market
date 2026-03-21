#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Append one MD test-case module to the combined Word document.
Usage: python md_to_word.py <md_filename>
Example: python md_to_word.py 用户认证模块.md
"""
import sys
import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paths ──────────────────────────────────────────────────────────────────
INPUT_DIR  = r'e:\IdeaProjects\flower-market\docs\test'
OUTPUT_DOC = r'e:\IdeaProjects\flower-market\docs\test\Word\测试用例汇总.docx'

# ── Border helpers ─────────────────────────────────────────────────────────
def _border(val='single', sz='12', space='0', color='000000'):
    b = OxmlElement('w:border')  # placeholder – will be replaced per-edge
    return val, sz, space, color

def set_table_borders(table):
    """Apply single-line borders to every cell in *table*."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '8')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def shade_cell(cell, fill="E8E8E8"):
    """Apply a light grey background to a cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def set_col_width(cell, width_cm):
    """Set the preferred width of a cell."""
    tcPr = cell._element.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_cm * 567)))  # 1 cm ≈ 567 twips
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


# ── Text helpers ──────────────────────────────────────────────────────────
FONT_NAME = 'Times New Roman'
FONT_SIZE = Pt(12)  # 小四号


def set_doc_default_font(doc):
    """Set Times New Roman 12pt as the document-level default."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE
    # Also set the East-Asian / complex-script font via rPr XML
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'),    FONT_NAME)
    rFonts.set(qn('w:hAnsi'),   FONT_NAME)
    rFonts.set(qn('w:cs'),      FONT_NAME)


def add_runs_with_linebreaks(paragraph, text, font_size=None, bold=False):
    """Split on '<br />' and add runs with XML linebreaks in between."""
    if font_size is None:
        font_size = FONT_SIZE
    parts = re.split(r'\s*<br\s*/>\s*', text)
    for idx, part in enumerate(parts):
        if part == '':
            continue
        run = paragraph.add_run(part)
        run.font.name = FONT_NAME
        run.font.size = font_size
        run.font.bold = bold
        if idx < len(parts) - 1 and parts[idx + 1] != '':
            run._r.append(OxmlElement('w:br'))


# ── Markdown parser ────────────────────────────────────────────────────────
def parse_md(filepath):
    """
    Return a list of parsed blocks:
      {'type': 'h1'|'h2'|'h3', 'text': ...}
      {'type': 'paragraph', 'text': ...}
      {'type': 'kv_table', 'rows': [(key, value), ...]}   ← test‑case table
      {'type': 'grid_table', 'headers': [...], 'rows': [[...]]}  ← coverage table
    """
    with open(filepath, encoding='utf-8') as f:
        lines = f.readlines()

    blocks = []
    i = 0

    while i < len(lines):
        line = lines[i].rstrip('\n')
        stripped = line.strip()

        # Headings
        m = re.match(r'^(#{1,3})\s+(.*)', stripped)
        if m:
            level = len(m.group(1))
            blocks.append({'type': f'h{level}', 'text': m.group(2).strip()})
            i += 1
            continue

        # Horizontal rule
        if stripped in ('---', '***', '___'):
            blocks.append({'type': 'hr'})
            i += 1
            continue

        # Table
        if stripped.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].rstrip('\n').strip())
                i += 1

            # Parse rows (skip separator rows made of --)
            parsed_rows = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.split('|')[1:-1]]
                # Skip pure separator rows
                if all(re.match(r'^[-:]+$', c) for c in cells if c):
                    continue
                parsed_rows.append(cells)

            if not parsed_rows:
                continue

            # Detect table type:
            # KV table = 2 columns where the first column looks like a field label
            if (len(parsed_rows[0]) == 2
                    and re.search(r'Test Case|Title|Module|Precond|Input|Step|Result|Status|Priority',
                                  parsed_rows[0][0], re.I)):
                blocks.append({'type': 'kv_table',
                                'rows': [(r[0], r[1]) for r in parsed_rows]})
            else:
                # Normal grid table: first row is header
                headers = parsed_rows[0]
                data    = parsed_rows[1:]
                blocks.append({'type': 'grid_table',
                                'headers': headers,
                                'rows': data})
            continue

        # Blank line
        if not stripped:
            i += 1
            continue

        # Regular paragraph
        blocks.append({'type': 'paragraph', 'text': stripped})
        i += 1

    return blocks


# ── Word builder ───────────────────────────────────────────────────────────
def append_module(doc, blocks, module_name):
    """Append parsed blocks to *doc* as a new module section."""
    # ── Page break before each module (except the very first paragraph)
    if len(doc.paragraphs) > 1 or any(t.text for t in doc.paragraphs):
        doc.add_page_break()

    for block in blocks:
        btype = block['type']

        # ── Headings ──────────────────────────────────────────────────────
        if btype == 'h1':
            h = doc.add_heading(block['text'], level=1)
            h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in h.runs:
                run.font.name = FONT_NAME
                run.font.size = Pt(16)

        elif btype == 'h2':
            h = doc.add_heading(block['text'], level=2)
            for run in h.runs:
                run.font.name = FONT_NAME
                run.font.size = Pt(14)

        elif btype == 'h3':
            h = doc.add_heading(block['text'], level=3)
            for run in h.runs:
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE

        # ── Paragraph ─────────────────────────────────────────────────────
        elif btype == 'paragraph':
            p = doc.add_paragraph()
            add_runs_with_linebreaks(p, block['text'])

        # ── HR ────────────────────────────────────────────────────────────
        elif btype == 'hr':
            pass  # omit decorative rules

        # ── KV test‑case table ────────────────────────────────────────────
        elif btype == 'kv_table':
            tbl = doc.add_table(rows=len(block['rows']), cols=2)
            tbl.style = 'Table Grid'

            # Column widths: label ~4.5 cm, value ~12.5 cm  (total ~17 cm in landscape)
            for row_idx, (key, val) in enumerate(block['rows']):
                cells = tbl.rows[row_idx].cells
                key_cell = cells[0]
                val_cell = cells[1]

                # Set column widths
                set_col_width(key_cell, 4.5)
                set_col_width(val_cell, 12.5)

                # Key cell – bold, no background
                key_cell.text = ''
                kp = key_cell.paragraphs[0]
                kp.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                add_runs_with_linebreaks(kp, key, bold=True)

                # Value cell
                val_cell.text = ''
                vp = val_cell.paragraphs[0]
                add_runs_with_linebreaks(vp, val)

            set_table_borders(tbl)
            # Each test-case table starts a new page (page break AFTER the table)
            doc.add_page_break()

        # ── Grid table (e.g. coverage summary) ───────────────────────────
        elif btype == 'grid_table':
            headers = block['headers']
            rows    = block['rows']
            if not headers:
                continue

            tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
            tbl.style = 'Table Grid'

            # Header row
            hdr_cells = tbl.rows[0].cells
            for ci, h in enumerate(headers):
                hdr_cells[ci].text = ''
                hp = hdr_cells[ci].paragraphs[0]
                hp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                add_runs_with_linebreaks(hp, h, bold=True)

            # Data rows
            for ri, row in enumerate(rows):
                dcells = tbl.rows[ri + 1].cells
                for ci, cell_text in enumerate(row):
                    if ci < len(dcells):
                        dcells[ci].text = ''
                        dp = dcells[ci].paragraphs[0]
                        add_runs_with_linebreaks(dp, cell_text)

            set_table_borders(tbl)
            doc.add_paragraph()


# ── Main ──────────────────────────────────────────────────────────────────
def main():
    if len(sys.argv) < 2:
        print("Usage: python md_to_word.py <md_filename>")
        print("Example: python md_to_word.py 用户认证模块.md")
        sys.exit(1)

    md_filename = sys.argv[1]
    md_path = os.path.join(INPUT_DIR, md_filename)

    if not os.path.exists(md_path):
        print(f"ERROR: file not found – {md_path}")
        sys.exit(1)

    # Open or create the combined document
    os.makedirs(os.path.dirname(OUTPUT_DOC), exist_ok=True)

    if os.path.exists(OUTPUT_DOC):
        doc = Document(OUTPUT_DOC)
        print(f"Opening existing document: {OUTPUT_DOC}")
    else:
        doc = Document()
        # Set default font: Times New Roman 小四号 (12pt)
        set_doc_default_font(doc)
        # Page setup: A4, landscape for wide tables
        section = doc.sections[0]
        section.page_width  = Cm(29.7)
        section.page_height = Cm(21.0)
        section.left_margin  = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin   = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        print(f"Creating new document: {OUTPUT_DOC}")

    # Parse and append
    blocks = parse_md(md_path)
    append_module(doc, blocks, md_filename)

    # Save
    doc.save(OUTPUT_DOC)
    print(f"OK – appended '{md_filename}' → {OUTPUT_DOC}")


if __name__ == '__main__':
    main()
